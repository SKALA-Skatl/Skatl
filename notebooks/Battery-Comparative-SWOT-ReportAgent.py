# Auto-generated from 06-Battery-Comparative-SWOT-ReportAgent.ipynb
# Generated for collaboration and Git-friendly review.

# %% [markdown]
# # Battery Comparative SWOT + Report Agent
#
# 이 노트북은 아키텍처의 마지막 단계를 담당합니다.
#
# - 입력: `SK On Agent`, `CATL Agent`, `Human Review #2` 결과 JSON
# - 처리: `Comparative SWOT Agent` -> `Report Agent`
# - 검토: `Human Review #3`를 위해 저장 직전 interrupt
# - 출력: 최종 Word 보고서(.docx)
#
# 초보자용 흐름으로 구성했기 때문에, 각 단계가 눈에 보이도록 상태와 출력 계약을 명시적으로 나눴습니다.

# %% [markdown]
# ## 1. 준비
#
# 이 노트북은 앞단 에이전트가 이미 JSON을 만들어줬다고 가정하고, 마지막 단계만 mock 데이터로 재현합니다.

# %%
from dotenv import load_dotenv
from langsmith import Client

load_dotenv(override=True)
langsmith_client = Client()

# %%
import json
from pathlib import Path
from typing import Dict, List, Literal, TypedDict

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import END, START, StateGraph
from langgraph.types import Command, interrupt
from pydantic import BaseModel, Field

# %% [markdown]
# ## 2. 입력 계약
#
# 내 역할은 앞단 출력을 받아 비교 분석을 수행하는 것이므로, 입력 JSON 스키마를 먼저 고정합니다.

# %%
class FinalStageState(TypedDict, total=False):
    market_context_payload: Dict
    reference_catalog_payload: Dict
    skon_payload: Dict
    catl_payload: Dict
    review2_payload: Dict
    comparative_swot: Dict
    report_draft: Dict
    human_decision: str
    human_feedback: str
    review_round: int
    max_review_rounds: int
    review_history: List[Dict]
    final_revision_mode: bool
    final_status: str
    final_report_path: str

# %% [markdown]
# ## 3. 동일 기준 SWOT 축
#
# 두 회사를 같은 잣대로 보기 위해 비교 기준을 명시적으로 고정합니다.

# %%
# Market Agent output 기반 비교 축
COMPARISON_AXES = [
    "chasm_response_fit",
    "market_position_and_scale",
    "technology_trend_fit",
    "ess_hev_diversification_fit",
    "policy_and_regulatory_resilience",
    "cost_competitiveness_under_price_decline",
]

MAX_HUMAN_REVIEW_ROUNDS = 3

# %% [markdown]
# ## 4. 출력 스키마
#
# 마지막 단계에서는 자유 텍스트보다 `비교 SWOT JSON`과 `리포트 JSON`을 먼저 만들고, 마지막에 Word(.docx)로 저장합니다.

# %%
class SWOTQuadrant(BaseModel):
    company: str = Field(description="회사명")
    strengths: List[str] = Field(description="강점")
    weaknesses: List[str] = Field(description="약점")
    opportunities: List[str] = Field(description="기회")
    threats: List[str] = Field(description="위협")


class ComparativeAxisResult(BaseModel):
    axis: str = Field(description="Market Agent가 포착한 시장 변화에 대한 대응 적합성을 비교하는 축")
    winner: Literal["SK On", "CATL", "tie"] = Field(description="해당 시장 변화에 더 적절히 대응한다고 판단되는 기업")
    rationale: str = Field(description="시장 데이터와 기업 전략을 연결한 판단 근거")


class SWOTComparisonTableRow(BaseModel):
    category: Literal["강점(S) - 내부 경쟁력", "약점(W) - 내부 취약점", "기회(O) - 외부 시장", "위협(T) - 외부 리스크"] = Field(description="SWOT 구분")
    company_a_summary: str = Field(description="A기업(SK On) 요약")
    company_b_summary: str = Field(description="B기업(CATL) 요약")
    strategic_implication: str = Field(description="전략적 시사점 또는 비교 평가")


class CompanySection(BaseModel):
    company: str = Field(description="회사명")
    portfolio_diversification: List[str] = Field(description="포트폴리오 다각화 내용")
    core_competencies: List[str] = Field(description="핵심 경쟁력")
    strategic_direction: str = Field(description="현재 전략 방향")
    key_watchpoints: List[str] = Field(description="추가로 볼 포인트")


class ComparativeSWOTOutput(BaseModel):
    comparison_axes: List[ComparativeAxisResult] = Field(description="Market Agent 기준의 동일 기준 비교 결과")
    swot_focus_points: List[str] = Field(description="시장 변화 대응 관점에서 본 S/W/O/T 비교 기준 및 주목 포인트")
    strategic_interactions: List[str] = Field(description="한 기업의 시장 대응 강점이 상대 기업의 위협·약점으로 연결되는 전략적 상호작용")
    skon_swot: SWOTQuadrant = Field(description="SK On SWOT")
    catl_swot: SWOTQuadrant = Field(description="CATL SWOT")
    swot_comparison_table: List[SWOTComparisonTableRow] = Field(description="SWOT 비교 표용 4개 행")
    strategic_gaps: List[str] = Field(description="시장 변화 대응 관점에서 드러나는 양사 전략 차이")
    decision_takeaways: List[str] = Field(description="의사결정자 관점에서 봐야 할 핵심 대응 포인트")
    evidence_used: List[str] = Field(description="사용한 evidence id 목록")
    confidence: float = Field(description="0과 1 사이 신뢰도")


class FinalReportOutput(BaseModel):
    title: str = Field(description="보고서 제목")
    summary: str = Field(description="기업 자체의 절대 평가보다 시장 변화 대응 적합성에 초점을 둔 핵심 메시지 요약. 반 페이지를 넘지 않는 분량")
    market_background: List[str] = Field(description="Market Agent가 포착한 배터리 시장 환경 변화와 전략적 의미")
    sk_on_section: CompanySection = Field(description="SK On이 시장 변화에 어떻게 대응하고 있는지 보여주는 포트폴리오 다각화 및 핵심 경쟁력")
    catl_section: CompanySection = Field(description="CATL이 시장 변화에 어떻게 대응하고 있는지 보여주는 포트폴리오 다각화 및 핵심 경쟁력")
    comparative_swot_focus_points: List[str] = Field(description="시장 변화 대응 관점에서 본 SWOT 비교 기준 및 주목 포인트")
    comparative_swot_company_comparison: List[str] = Field(description="같은 시장 변화에 대한 SK On과 CATL의 대응 차이를 보여주는 S/W/O/T 별 기업 비교")
    integrated_implications: List[str] = Field(description="현재 우위 기업, 우위가 뒤집힐 조건, SK On 우선 과제, CATL 리스크 모니터링 포인트, 의사결정자가 볼 핵심 신호를 포함한 종합 시사점")
    references: List[str] = Field(description="실제 활용 자료 목록. 지정 형식 준수")


swot_parser = JsonOutputParser(pydantic_object=ComparativeSWOTOutput)
report_parser = JsonOutputParser(pydantic_object=FinalReportOutput)

# %% [markdown]
# ## 5. LLM

# %%
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)

# %% [markdown]
# ## 6. Comparative SWOT Agent

# %%
def comparative_swot_agent(state: FinalStageState):
    prompt = PromptTemplate(
        template="""
        당신은 배터리 시장 전략 비교 분석가입니다.
        입력으로 주어진 JSON만 사용해서 SK On과 CATL의 Comparative SWOT를 작성하세요.

        목표:
        - 동일 기준 기반 Comparative SWOT 비교
        - 기업 자체의 절대적 우열보다, Market Agent가 포착한 시장 변화에 각 기업이 얼마나 잘 맞게 대응하고 있는지 비교
        - 전략적 우위 및 리스크 요인 도출
        - A기업 강점이 B기업 위협 또는 약점으로 이어지는 전략적 상호작용을 한눈에 보이게 정리
        - 의사결정자가 바로 사용할 수 있는 비교 포인트 정리

        규칙:
        - 외부 지식을 추가하지 말 것
        - evidence id를 반드시 evidence_used에 남길 것
        - comparison_axes는 아래 축만 사용할 것
        - winner는 SK On, CATL, tie 중 하나만 사용
        - swot_focus_points는 S/W/O/T 각각의 비교 기준과 주목 포인트가 드러나게 작성할 것
        - 모든 서술형 문장과 bullet, 표 셀 내용은 반드시 한국어로 작성할 것
        - 회사명, 배터리 화학계열(LFP, NCM), 제도명(IRA), 단위, 고유명사 외에는 영어 문장을 쓰지 말 것
        - 분석의 중심 질문은 "누가 더 좋은 기업인가"가 아니라 "누가 현재 시장 변화에 더 적절히 대응하고 있는가"여야 함
        - 반드시 아래 MarketContext 6개 섹션을 해석에 활용할 것: ev_growth_slowdown, market_share_ranking, lfp_ncm_trend, ess_hev_growth, regulatory_status, cost_competitiveness
        - 각 comparison axis는 아래 의미를 따른다
          - chasm_response_fit: EV 성장 둔화와 지역별 캐즘에 대한 대응 적합성
          - market_position_and_scale: 글로벌 점유율과 체급 기반 방어력
          - technology_trend_fit: LFP vs NCM 트렌드와 기술 포트폴리오의 정합성
          - ess_hev_diversification_fit: ESS/HEV 성장 기회와 다각화 전략의 적합성
          - policy_and_regulatory_resilience: IRA, 관세, EU 규제에 대한 대응력과 리스크
          - cost_competitiveness_under_price_decline: 배터리 가격 하락 국면에서의 원가 방어력
        - strategic_interactions는 최소 3개 작성하고, 한 회사의 강점이 다른 회사의 위협 또는 약점으로 연결되는 방식으로 쓸 것
        - swot_comparison_table은 아래 4개 category를 정확히 한 번씩 포함할 것
        - review2_payload의 지시사항을 반영할 것
        - 반드시 JSON으로만 답할 것

        comparison_axes:
        {comparison_axes}

        swot_table_categories:
        ["강점(S) - 내부 경쟁력", "약점(W) - 내부 취약점", "기회(O) - 외부 시장", "위협(T) - 외부 리스크"]

        market_context_payload:
        {market_context_payload}

        review2_payload:
        {review2_payload}

        human_feedback_from_previous_review:
        {human_feedback}

        final_revision_mode:
        {final_revision_mode}

        SK On payload:
        {skon_payload}

        CATL payload:
        {catl_payload}

        {format_instructions}
        """,
        input_variables=["comparison_axes", "market_context_payload", "review2_payload", "human_feedback", "final_revision_mode", "skon_payload", "catl_payload"],
        partial_variables={"format_instructions": swot_parser.get_format_instructions()},
    )

    chain = prompt | llm | swot_parser
    result = chain.invoke(
        {
            "comparison_axes": COMPARISON_AXES,
            "market_context_payload": json.dumps(state["market_context_payload"], ensure_ascii=False, indent=2),
            "review2_payload": json.dumps(state["review2_payload"], ensure_ascii=False, indent=2),
            "human_feedback": state.get("human_feedback", "No prior human feedback."),
            "final_revision_mode": state.get("final_revision_mode", False),
            "skon_payload": json.dumps(state["skon_payload"], ensure_ascii=False, indent=2),
            "catl_payload": json.dumps(state["catl_payload"], ensure_ascii=False, indent=2),
        }
    )
    return {"comparative_swot": result}

# %% [markdown]
# ## 7. Report Agent

# %%
def report_agent(state: FinalStageState):
    market_refs = [
        url
        for url in state["reference_catalog_payload"].get("references", {})
        if "/market/" in url
    ]
    skon_refs = [item["source_url"] for item in state["skon_payload"].get("evidence", [])]
    catl_refs = [item["source_url"] for item in state["catl_payload"].get("evidence", [])]
    all_ref_urls = list(dict.fromkeys(market_refs + skon_refs + catl_refs))
    reference_catalog = state["reference_catalog_payload"].get("references", {})
    formatted_references = [reference_catalog[url] for url in all_ref_urls if url in reference_catalog]

    prompt = PromptTemplate(
        template="""
        당신은 의사결정자용 배터리 시장 전략 보고서 작성자입니다.
        아래 입력만 사용해 보고서 초안을 작성하세요.

        보고서 목차:
        1. Summary
        2. 시장 배경(배터리 시장 환경 변화)
        3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력
           - SK ON
           - CATL
        4. Comparative SWOT
           - S,W,O,T 별 비교 기준 및 주목한 포인트
           - S,W,O,T 별 기업 비교
           - SWOT 비교 요약 table
        5. 종합 시사점
        6. Reference

        작성 규칙:
        - Summary는 보고서 전체 핵심 메시지 중심으로 작성하고, 1/2 페이지를 넘지 않게 압축할 것
        - summary를 포함한 모든 서술형 문장, bullet, 표 설명은 반드시 한국어로 작성할 것
        - 회사명, 배터리 화학계열(LFP, NCM), 제도명(IRA), 단위, 고유명사 외에는 영어 문장을 쓰지 말 것
        - 보고서 전체 논리는 "기업 자체가 얼마나 좋은가"보다 "Market Agent가 포착한 시장 변화에 얼마나 잘 맞게 대응하고 있는가" 중심으로 전개할 것
        - 시장 배경은 MarketContext의 6개 섹션이 보여주는 시장 변화와 전략적 의미를 bullet로 정리할 것
        - 기업별 섹션은 MarketContext를 기준으로 각 기업의 포트폴리오 다각화와 핵심 경쟁력이 시장 변화에 얼마나 맞는지 쓰고, 포트폴리오 다각화, 핵심 경쟁력, 전략 방향, watchpoint를 분리해서 작성할 것
        - Comparative SWOT에서는 swot_focus_points, strategic_interactions, swot_comparison_table을 모두 반영할 것
        - Comparative SWOT은 다음 질문에 답하는 방향이어야 함: 캐즘 대응력, 체급과 점유율, 기술 트렌드 적합성, ESS/HEV 기회 활용, 규제 대응력, 가격 하락 국면 원가 방어력
        - A기업의 강점이 B기업의 위협이나 약점으로 직결되는 전략적 상호작용이 드러나야 함
        - 종합 시사점은 최소 4개 이상 작성할 것
        - 종합 시사점에는 현재 기준 우위 기업, 우위가 뒤집힐 조건, SK On의 우선 대응 과제, CATL의 핵심 리스크 또는 모니터링 포인트를 포함할 것
        - references는 제공된 형식화된 목록만 사용하고, 실제 활용한 자료만 남길 것
        - 반드시 JSON으로만 답할 것

        Market context:
        {market_context_payload}

        Comparative SWOT:
        {comparative_swot}

        SK On payload:
        {skon_payload}

        CATL payload:
        {catl_payload}

        Allowed formatted references:
        {references}

        human_feedback_from_previous_review:
        {human_feedback}

        final_revision_mode:
        {final_revision_mode}

        {format_instructions}
        """,
        input_variables=["market_context_payload", "comparative_swot", "skon_payload", "catl_payload", "references", "human_feedback", "final_revision_mode"],
        partial_variables={"format_instructions": report_parser.get_format_instructions()},
    )

    chain = prompt | llm | report_parser
    result = chain.invoke(
        {
            "market_context_payload": json.dumps(state["market_context_payload"], ensure_ascii=False, indent=2),
            "comparative_swot": json.dumps(state["comparative_swot"], ensure_ascii=False, indent=2),
            "skon_payload": json.dumps(state["skon_payload"], ensure_ascii=False, indent=2),
            "catl_payload": json.dumps(state["catl_payload"], ensure_ascii=False, indent=2),
            "references": json.dumps(formatted_references, ensure_ascii=False, indent=2),
            "human_feedback": state.get("human_feedback", "No prior human feedback."),
            "final_revision_mode": state.get("final_revision_mode", False),
        }
    )
    return {"report_draft": result}

# %% [markdown]
# ## 8. 최종 저장 노드
#
# 이 노드는 Human Review #3 이후 실행된다고 가정합니다.

# %%
ACCENT_COLOR = RGBColor(0xFB, 0x77, 0x62)
SEPARATOR_COLOR = "D9D9D9"
TITLE_SEPARATOR_COLOR = "FB7762"
TABLE_HEADER_FILL = "FDE3DE"
TABLE_FIRST_COL_FILL = "FEF1EE"


def _add_bullet_items(doc: Document, items: List[str]):
    for item in items:
        paragraph = doc.add_paragraph()
        _remove_numbering(paragraph)
        paragraph.add_run(f"· {item}")


def _remove_numbering(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def _add_clean_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    _remove_numbering(paragraph)
    run = paragraph.add_run(text)
    run.font.color.rgb = ACCENT_COLOR
    return paragraph


def _set_run_size(run, size_pt: int):
    run.font.size = Pt(size_pt)


def _add_title(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    _remove_numbering(paragraph)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = ACCENT_COLOR
    _set_run_size(run, 22)
    return paragraph


def _add_separator_line(doc: Document, color: str = SEPARATOR_COLOR, val: str = "single", size: str = "6"):
    paragraph = doc.add_paragraph()
    _remove_numbering(paragraph)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), val)
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_company_section(doc: Document, heading: str, section: Dict):
    _add_clean_heading(doc, heading, level=2)
    _add_clean_heading(doc, "포트폴리오 다각화", level=3)
    _add_bullet_items(doc, section.get("portfolio_diversification", []))
    _add_clean_heading(doc, "핵심 경쟁력", level=3)
    _add_bullet_items(doc, section.get("core_competencies", []))
    _add_clean_heading(doc, "전략 방향", level=3)
    _add_bullet_items(doc, [section.get("strategic_direction", "")])
    watchpoints = section.get("key_watchpoints", [])
    if watchpoints:
        _add_clean_heading(doc, "Watchpoints", level=3)
        _add_bullet_items(doc, watchpoints)


def _set_cell_background(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_swot_comparison_table(doc: Document, swot: Dict):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    header_cells[0].text = "구분"
    header_cells[1].text = "A기업 (SK On)"
    header_cells[2].text = "B기업 (CATL)"
    header_cells[3].text = "전략적 시사점 (비교 평가)"
    for cell in header_cells:
        _set_cell_background(cell, TABLE_HEADER_FILL)

    for row in swot.get("swot_comparison_table", []):
        cells = table.add_row().cells
        cells[0].text = row["category"]
        cells[1].text = row["company_a_summary"]
        cells[2].text = row["company_b_summary"]
        cells[3].text = row["strategic_implication"]
        _set_cell_background(cells[0], TABLE_FIRST_COL_FILL)


def build_word_report(report: Dict, swot: Dict, output_path: Path) -> None:
    doc = Document()
    _add_title(doc, report["title"])
    _add_separator_line(doc, color=TITLE_SEPARATOR_COLOR, val="double", size="10")

    _add_clean_heading(doc, "목차", level=1)
    _add_bullet_items(
        doc,
        [
            "1. Summary",
            "2. 시장 배경(배터리 시장 환경 변화)",
            "3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력",
            "4. Comparative SWOT",
            "5. 종합 시사점",
            "6. Reference",
        ],
    )
    _add_separator_line(doc)

    _add_clean_heading(doc, "1. Summary", level=1)
    doc.add_paragraph(report["summary"])

    _add_clean_heading(doc, "2. 시장 배경(배터리 시장 환경 변화)", level=1)
    _add_bullet_items(doc, report["market_background"])

    _add_clean_heading(doc, "3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력", level=1)
    _add_company_section(doc, "3-1. SK On", report["sk_on_section"])
    _add_company_section(doc, "3-2. CATL", report["catl_section"])

    _add_clean_heading(doc, "4. Comparative SWOT", level=1)
    _add_clean_heading(doc, "4-1. S,W,O,T 별 비교 기준 및 주목한 포인트", level=2)
    _add_bullet_items(doc, report["comparative_swot_focus_points"])
    _add_clean_heading(doc, "4-2. S,W,O,T 별 기업 비교", level=2)
    _add_bullet_items(doc, report["comparative_swot_company_comparison"])
    _add_clean_heading(doc, "4-3. SWOT 비교 요약 table", level=2)
    _add_swot_comparison_table(doc, swot)

    _add_clean_heading(doc, "5. 종합 시사점", level=1)
    _add_bullet_items(doc, report["integrated_implications"])

    _add_clean_heading(doc, "6. Reference", level=1)
    _add_bullet_items(doc, report["references"])

    doc.save(output_path)


def _normalize_review_response(response) -> Dict[str, str]:
    if isinstance(response, dict) and "decisions" in response:
        decisions = response.get("decisions", [])
        if not decisions:
            return {"decision": "reject", "feedback": "No review decision provided."}

        first_decision = decisions[0]
        decision_type = str(first_decision.get("type", "reject")).strip().lower()
        if decision_type == "approve":
            return {"decision": "approve", "feedback": ""}
        if decision_type == "reject":
            return {"decision": "reject", "feedback": str(first_decision.get("message", "")).strip()}
        return {"decision": "reject", "feedback": f"Unsupported decision type: {decision_type}"}

    if isinstance(response, dict):
        decision = str(response.get("decision", "reject")).strip().lower()
        feedback = str(response.get("feedback", "")).strip()
        if decision not in {"approve", "reject"}:
            decision = "reject"
        return {"decision": decision, "feedback": feedback}

    return {"decision": "reject", "feedback": "Invalid human review response."}


def human_review(state: FinalStageState):
    next_round = state.get("review_round", 0) + 1
    response = interrupt(
        {
            "stage": "human_review_3",
            "review_round": next_round,
            "max_review_rounds": state.get("max_review_rounds", MAX_HUMAN_REVIEW_ROUNDS),
            "action_requests": [
                {
                    "name": "review_report_draft",
                    "args": {
                        "report_draft": state["report_draft"],
                        "comparative_swot": state["comparative_swot"],
                    },
                    "description": "Review the comparative SWOT result and report draft before publication.",
                }
            ],
            "review_configs": [
                {
                    "action_name": "review_report_draft",
                    "allowed_decisions": ["approve", "reject"],
                }
            ],
        }
    )

    normalized = _normalize_review_response(response)
    history = list(state.get("review_history", []))
    history.append(
        {
            "review_round": next_round,
            "decision": normalized["decision"],
            "feedback": normalized["feedback"],
        }
    )

    return {
        "human_decision": normalized["decision"],
        "human_feedback": normalized["feedback"],
        "review_round": next_round,
        "review_history": history,
        "final_revision_mode": False,
    }


def format_human_review_preview(interrupt_data: Dict) -> str:
    action = interrupt_data["action_requests"][0]
    args = action["args"]
    report = args["report_draft"]
    swot = args["comparative_swot"]

    lines = [
        "# Human Review Preview",
        "",
        f"- Stage: {interrupt_data.get('stage', 'human_review_3')}",
        f"- Review Round: {interrupt_data.get('review_round')}/{interrupt_data.get('max_review_rounds')}",
        f"- Allowed Decisions: {', '.join(interrupt_data['review_configs'][0]['allowed_decisions'])}",
        "",
        "# 목차",
        "1. Summary",
        "2. 시장 배경(배터리 시장 환경 변화)",
        "3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력",
        "4. Comparative SWOT",
        "5. 종합 시사점",
        "6. Reference",
        "",
        "--------------------------------------------------",
        "",
        "# 1. Summary",
        report["summary"],
        "",
        "# 2. 시장 배경(배터리 시장 환경 변화)",
    ]
    lines.extend([f"- {item}" for item in report.get("market_background", [])])
    lines.extend([
        "",
        "# 3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력",
        "## 3-1. SK ON",
        "### 포트폴리오 다각화",
    ])
    lines.extend([f"  - {item}" for item in report.get("sk_on_section", {}).get("portfolio_diversification", [])])
    lines.extend(["### 핵심 경쟁력"])
    lines.extend([f"  - {item}" for item in report.get("sk_on_section", {}).get("core_competencies", [])])
    lines.extend([
        "### 전략 방향",
    ])
    lines.extend([f"  - {report.get('sk_on_section', {}).get('strategic_direction', '')}"])
    lines.extend(["### Watchpoints"])
    lines.extend([f"  - {item}" for item in report.get("sk_on_section", {}).get("key_watchpoints", [])])
    lines.extend([
        "",
        "## 3-2. CATL",
        "### 포트폴리오 다각화",
    ])
    lines.extend([f"  - {item}" for item in report.get("catl_section", {}).get("portfolio_diversification", [])])
    lines.extend(["### 핵심 경쟁력"])
    lines.extend([f"  - {item}" for item in report.get("catl_section", {}).get("core_competencies", [])])
    lines.extend([
        "### 전략 방향",
    ])
    lines.extend([f"  - {report.get('catl_section', {}).get('strategic_direction', '')}"])
    lines.extend(["### Watchpoints"])
    lines.extend([f"  - {item}" for item in report.get("catl_section", {}).get("key_watchpoints", [])])
    lines.extend([
        "",
        "# 4. Comparative SWOT",
        "## 4-1. S,W,O,T 별 비교 기준 및 주목한 포인트",
    ])
    lines.extend([f"- {item}" for item in report.get("comparative_swot_focus_points", [])])
    lines.extend([
        "",
        "## 4-2. S,W,O,T 별 기업 비교",
    ])
    lines.extend([f"- {item}" for item in report.get("comparative_swot_company_comparison", [])])
    lines.extend([
        "",
        "## 4-3. SWOT 비교 요약 table",
        "| 구분 | A기업 (SK On) | B기업 (CATL) | 전략적 시사점 |",
        "|---|---|---|---|",
    ])
    for row in swot.get("swot_comparison_table", []):
        lines.append(
            f"| {row['category']} | {row['company_a_summary']} | {row['company_b_summary']} | {row['strategic_implication']} |"
        )
    lines.extend([
        "",
        "# 5. 종합 시사점",
    ])
    lines.extend([f"- {item}" for item in report.get("integrated_implications", [])])
    lines.extend([
        "",
        "# 6. Reference",
    ])
    lines.extend([f"- {item}" for item in report.get("references", [])])
    return "\n".join(lines)


def save_human_review_preview(interrupt_data: Dict, output_dir: Path = Path("./for_check")) -> Dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    round_no = interrupt_data.get("review_round", "unknown")

    preview_md_path = output_dir / f"human_review_round_{round_no}_preview.md"
    preview_json_path = output_dir / f"human_review_round_{round_no}_payload.json"

    preview_md_path.write_text(format_human_review_preview(interrupt_data), encoding="utf-8")
    preview_json_path.write_text(
        json.dumps(interrupt_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return {
        "preview_markdown_path": str(preview_md_path),
        "preview_json_path": str(preview_json_path),
    }


def route_after_human_review(state: FinalStageState) -> str:
    decision = state.get("human_decision", "reject")
    review_round = state.get("review_round", 0)
    max_rounds = state.get("max_review_rounds", MAX_HUMAN_REVIEW_ROUNDS)

    if decision == "approve":
        return "publish_report"
    if review_round >= max_rounds:
        return "final_revision"
    return "comparative_swot_agent"


def enable_final_revision(state: FinalStageState):
    return {
        "final_revision_mode": True,
        "final_status": "auto_publish_after_max_rejects",
    }


def publish_report(state: FinalStageState):
    output_dir = Path("./results")
    output_dir.mkdir(parents=True, exist_ok=True)

    report_path = output_dir / "battery_market_strategy_report_3.docx"
    build_word_report(state["report_draft"], state["comparative_swot"], report_path)

    return {
        "final_status": state.get("final_status", "approved"),
        "final_report_path": str(report_path),
    }

# %% [markdown]
# ## 9. Graph

# %%
builder = StateGraph(FinalStageState)
builder.add_node("comparative_swot_agent", comparative_swot_agent)
builder.add_node("report_agent", report_agent)
builder.add_node("human_review", human_review)
builder.add_node("publish_report", publish_report)
builder.add_node("final_revision", enable_final_revision)

builder.add_edge(START, "comparative_swot_agent")
builder.add_edge("comparative_swot_agent", "report_agent")
builder.add_edge("report_agent", "human_review")
builder.add_conditional_edges(
    "human_review",
    route_after_human_review,
    {
        "publish_report": "publish_report",
        "comparative_swot_agent": "comparative_swot_agent",
        "final_revision": "final_revision",
    },
)
builder.add_edge("final_revision", "comparative_swot_agent")
builder.add_edge("publish_report", END)

memory = MemorySaver()
graph = builder.compile(checkpointer=memory)

# %% [markdown]
# ## 10. Mock 입력 로드

# %%
base_dir = Path("../mock_data")

market_context_payload = json.loads((base_dir / "battery_market_background_payload.json").read_text(encoding="utf-8"))
reference_catalog_payload = json.loads((base_dir / "battery_reference_catalog.json").read_text(encoding="utf-8"))
skon_payload = json.loads((base_dir / "battery_strategy_skon_payload.json").read_text(encoding="utf-8"))
catl_payload = json.loads((base_dir / "battery_strategy_catl_payload.json").read_text(encoding="utf-8"))
review2_payload = json.loads((base_dir / "battery_strategy_human_review_2.json").read_text(encoding="utf-8"))

initial_state = {
    "market_context_payload": market_context_payload,
    "reference_catalog_payload": reference_catalog_payload,
    "skon_payload": skon_payload,
    "catl_payload": catl_payload,
    "review2_payload": review2_payload,
    "human_feedback": "",
    "review_round": 0,
    "max_review_rounds": MAX_HUMAN_REVIEW_ROUNDS,
    "review_history": [],
    "final_revision_mode": False,
}

# %% [markdown]
# ## 11. 첫 실행: Human Review #3에서 interrupt

# %%
config = {"configurable": {"thread_id": "battery-final-stage-001"}}

result = None
for event in graph.stream(initial_state, config=config, stream_mode="updates"):
    print(f"[executed nodes] {', '.join(event.keys())}")
    if "__interrupt__" in event:
        result = event

if result and "__interrupt__" in result:
    print("Human review interrupt received. Run the next cell for a readable preview.")
else:
    print(graph.get_state(config).values)

# %%
if "__interrupt__" in result:
    interrupt_data = result["__interrupt__"][0].value
    print("=== Interrupt Detected ===")
    print(interrupt_data["action_requests"][0]["name"])
    print(interrupt_data["review_configs"][0]["allowed_decisions"])
    print()
    print(format_human_review_preview(interrupt_data))
    print()
    saved_paths = save_human_review_preview(interrupt_data)
    print(saved_paths)
else:
    print("No interrupt occurred")

# %% [markdown]
# ## 12. 거절 후 재실행 예시
#
# 참고 노트북과 같은 방식으로 `Command(resume={"decisions": [...]})`를 사용합니다.
# 거절하면 feedback가 state에 저장되고, Comparative SWOT Agent부터 다시 실행됩니다.

# %%
result = None
for event in graph.stream(
    Command(
        resume={
            "decisions": [
                {
                    "type": "reject",
                    "message": "정책 리스크 비교를 더 명확히 써주세요.",
                }
            ]
        }
    ),
    config=config,
    stream_mode="updates",
):
    print(f"[executed nodes] {', '.join(event.keys())}")
    if "__interrupt__" in event:
        result = event

print(graph.get_state(config).values["review_history"])

if result and "__interrupt__" in result:
    interrupt_data = result["__interrupt__"][0].value
    print()
    print(format_human_review_preview(interrupt_data))
    print()
    saved_paths = save_human_review_preview(interrupt_data)
    print(saved_paths)

# %% [markdown]
# ## 13. 승인 후 보고서 생성 예시

# %%
result = None
for event in graph.stream(
    Command(resume={"decisions": [{"type": "approve"}]}),
    config=config,
    stream_mode="updates",
):
    print(f"[executed nodes] {', '.join(event.keys())}")
    if "__interrupt__" in event:
        result = event

final_state = graph.get_state(config)
print(final_state.values.get("final_status"), final_state.values.get("final_report_path"))

if result and "__interrupt__" in result:
    interrupt_data = result["__interrupt__"][0].value
    print()
    print(format_human_review_preview(interrupt_data))
    print()
    saved_paths = save_human_review_preview(interrupt_data)
    print(saved_paths)

# %% [markdown]
# 3번째 reject가 발생하면 `final_revision` 모드가 켜지고, Comparative SWOT Agent와 Report Agent를 한 번 더 거친 뒤 자동으로 Word 보고서를 생성합니다.

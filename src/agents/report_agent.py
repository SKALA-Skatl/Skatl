"""
Report Agent — Phase 4.

배터리 시장 전략 분석 최종 보고서를 생성하고 Word(.docx)로 저장한다.
노트북(Battery-Comparative-SWOT-ReportAgent.py)의 report_agent와
publish_report 로직을 모듈화한 버전이다.

입력:
  - market_context     : MarketContext (Market Agent 출력)
  - skon_result        : StrategyAgentOutput
  - catl_result        : StrategyAgentOutput
  - comparative_swot   : ComparativeSWOTOutput dict
  - human_feedback     : HITL #3 거절 피드백 (재실행 시)
  - final_revision_mode: 최대 리뷰 횟수 도달 시 최종 수정 모드 플래그

출력:
  FinalReportOutput (dict)

Word 저장:
  build_word_report(report_draft, comparative_swot, output_path)
  ※ python-docx 패키지 필요: pip install python-docx
"""

from __future__ import annotations
import json
from pathlib import Path
from typing import Dict, List

from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from pydantic import BaseModel, Field

from schemas.agent_io import SourceRecord, StrategyAgentOutput
from schemas.market_context import MarketContext
from logging_utils import get_logger


logger = get_logger("report_agent")


# ─────────────────────────────────────────────
# 출력 스키마
# ─────────────────────────────────────────────

class CompanySection(BaseModel):
    company: str = Field(description="회사명")
    portfolio_diversification: List[str] = Field(description="포트폴리오 다각화 내용")
    core_competencies: List[str] = Field(description="핵심 경쟁력")
    strategic_direction: str = Field(description="현재 전략 방향")
    key_watchpoints: List[str] = Field(description="추가로 볼 포인트")


class FinalReportOutput(BaseModel):
    title: str = Field(description="보고서 제목")
    summary: str = Field(
        description="기업 자체의 절대 평가보다 시장 변화 대응 적합성에 초점을 둔 핵심 메시지 요약. 반 페이지를 넘지 않는 분량"
    )
    market_background: List[str] = Field(
        description="Market Agent가 포착한 배터리 시장 환경 변화와 전략적 의미"
    )
    sk_on_section: CompanySection = Field(
        description="SK On이 시장 변화에 어떻게 대응하고 있는지 보여주는 포트폴리오 다각화 및 핵심 경쟁력"
    )
    catl_section: CompanySection = Field(
        description="CATL이 시장 변화에 어떻게 대응하고 있는지 보여주는 포트폴리오 다각화 및 핵심 경쟁력"
    )
    comparative_swot_focus_points: List[str] = Field(
        description="시장 변화 대응 관점에서 본 SWOT 비교 기준 및 주목 포인트"
    )
    comparative_swot_company_comparison: List[str] = Field(
        description="같은 시장 변화에 대한 SK On과 CATL의 대응 차이를 보여주는 S/W/O/T 별 기업 비교"
    )
    integrated_implications: List[str] = Field(
        description="어느 기업이 어떤 조건에서 더 유리한지와 우위가 뒤집힐 조건을 포함한 종합 시사점"
    )
    references: List[str] = Field(description="실제 활용 자료 목록. 지정 형식 준수")


_REPORT_PROMPT = """
당신은 의사결정자용 배터리 시장 전략 보고서 작성자입니다.
아래 입력만 사용해 보고서 초안을 작성하세요.

보고서 목차:
1. Summary
2. 시장 배경(배터리 시장 환경 변화)
3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력
   - SK ON
   - CATL
4. Comparative SWOT
   - S,W,O,T 별 비교 기준 및 주목한 포인트
   - S,W,O,T 별 기업 비교
   - SWOT 비교 요약 table
5. 종합 시사점
6. Reference

작성 규칙:
- Summary는 보고서 전체 핵심 메시지 중심으로 작성하고, 1/2 페이지를 넘지 않게 압축할 것
- summary를 포함한 모든 서술형 문장, bullet, 표 설명은 반드시 한국어로 작성할 것
- 회사명, 배터리 화학계열(LFP, NCM), 제도명(IRA), 단위, 고유명사 외에는 영어 문장을 쓰지 말 것
- 보고서 전체 논리는 "기업 자체가 얼마나 좋은가"보다 "Market Agent가 포착한 시장 변화에 얼마나 잘 맞게 대응하고 있는가" 중심으로 전개할 것
- 시장 배경은 MarketContext의 6개 섹션이 보여주는 시장 변화와 전략적 의미를 bullet로 정리할 것
- 기업별 섹션은 MarketContext를 기준으로 각 기업의 포트폴리오 다각화와 핵심 경쟁력이 시장 변화에 얼마나 맞는지 쓰고, 포트폴리오 다각화, 핵심 경쟁력, 전략 방향, watchpoint를 분리해서 작성할 것
- Comparative SWOT에서는 swot_focus_points, strategic_interactions, swot_comparison_table을 모두 반영할 것
- Comparative SWOT은 다음 질문에 답하는 방향이어야 함: 캐즘 대응력, 체급과 점유율, 기술 트렌드 적합성, ESS/HEV 기회 활용, 규제 대응력, 가격 하락 국면 원가 방어력
- A기업의 강점이 B기업의 위협이나 약점으로 직결되는 전략적 상호작용이 드러나야 함
- 종합 시사점은 향후 대응 방향과 우위가 뒤집힐 조건까지 포함할 것
- references는 제공된 형식화된 목록만 사용하고, 실제 활용한 자료만 남길 것
- human_feedback가 있으면 반드시 반영할 것
- final_revision_mode가 true이면 최대한 완성도 높게 작성할 것
- 반드시 JSON으로만 답할 것

Market context:
{market_context}

Comparative SWOT:
{comparative_swot}

SK On 전략 분석:
{skon_result}

CATL 전략 분석:
{catl_result}

Allowed references:
{references}

human_feedback_from_previous_review:
{human_feedback}

final_revision_mode:
{final_revision_mode}

{format_instructions}
"""


# ─────────────────────────────────────────────
# 참고문헌 빌더
# ─────────────────────────────────────────────

def _build_references(
    skon_result: StrategyAgentOutput,
    catl_result: StrategyAgentOutput,
) -> List[str]:
    """skon/catl 출처 목록에서 중복 제거한 참고문헌 문자열 리스트 반환."""
    seen: set[str] = set()
    refs: list[str] = []

    all_sources: list[SourceRecord] = (
        list(skon_result.get("sources", []))
        + list(catl_result.get("sources", []))
    )
    for src in all_sources:
        url = src.get("url", "")
        if url and url not in seen:
            seen.add(url)
            title = src.get("title", url)
            retrieved = src.get("retrieved_at", "")[:10]  # YYYY-MM-DD
            refs.append(f"{title}. Retrieved {retrieved}. {url}")

    return refs or ["내부 리서치 문서 및 공개 출처 참조"]


# ─────────────────────────────────────────────
# 실행 함수
# ─────────────────────────────────────────────

async def run_report_agent(
    market_context: MarketContext,
    skon_result: StrategyAgentOutput,
    catl_result: StrategyAgentOutput,
    comparative_swot: dict,
    human_feedback: str = "",
    final_revision_mode: bool = False,
) -> dict:
    """
    Report Agent 실행. 보고서 초안(JSON)을 생성한다.

    Args:
        market_context      : Market Agent 출력
        skon_result         : SKON Strategy Agent 출력
        catl_result         : CATL Strategy Agent 출력
        comparative_swot    : Comparative SWOT Agent 출력
        human_feedback      : HITL #3 거절 피드백 (재실행 시)
        final_revision_mode : 리뷰 한계 도달 시 최종 수정 모드

    Returns:
        FinalReportOutput을 dict로 변환한 결과
    """
    with logger.node_span("report_agent", {
        "has_feedback": bool(human_feedback),
        "final_revision": final_revision_mode,
    }):
        references = _build_references(skon_result, catl_result)

        parser = JsonOutputParser(pydantic_object=FinalReportOutput)
        prompt = PromptTemplate(
            template=_REPORT_PROMPT,
            input_variables=[
                "market_context", "comparative_swot",
                "skon_result", "catl_result",
                "references", "human_feedback", "final_revision_mode",
            ],
            partial_variables={"format_instructions": parser.get_format_instructions()},
        )
        llm   = ChatOpenAI(model="gpt-4o-mini", temperature=0)
        chain = prompt | llm | parser

        result = await chain.ainvoke({
            "market_context":      json.dumps(market_context,    ensure_ascii=False, indent=2),
            "comparative_swot":    json.dumps(comparative_swot,  ensure_ascii=False, indent=2),
            "skon_result":         json.dumps(skon_result,        ensure_ascii=False, indent=2),
            "catl_result":         json.dumps(catl_result,        ensure_ascii=False, indent=2),
            "references":          json.dumps(references,         ensure_ascii=False, indent=2),
            "human_feedback":      human_feedback or "No prior human feedback.",
            "final_revision_mode": final_revision_mode,
        })

        return result if isinstance(result, dict) else result.model_dump()


# ─────────────────────────────────────────────
# Word 문서 저장
# ─────────────────────────────────────────────

def build_word_report(
    report: Dict,
    swot: Dict,
    output_path: Path,
) -> None:
    """
    보고서 JSON + SWOT JSON → Word(.docx) 저장.

    Args:
        report      : FinalReportOutput dict
        swot        : ComparativeSWOTOutput dict
        output_path : 저장 경로 (.docx)

    Raises:
        ImportError : python-docx 미설치 시
    """
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import RGBColor
    except ImportError as e:
        raise ImportError(
            "Word 보고서 생성에는 python-docx가 필요합니다: pip install python-docx"
        ) from e

    ACCENT_COLOR       = RGBColor(0xFB, 0x77, 0x62)
    SEPARATOR_COLOR    = "FB7762"
    TABLE_HEADER_FILL  = "FDE3DE"
    TABLE_FIRST_COL_FILL = "FEF1EE"

    def _remove_numbering(paragraph):
        p_pr  = paragraph._p.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)

    def _add_bullet_items(doc: Document, items: List[str]):
        for item in items:
            p = doc.add_paragraph()
            _remove_numbering(p)
            p.add_run(f"· {item}")

    def _add_clean_heading(doc: Document, text: str, level: int = 1):
        p = doc.add_paragraph(style=f"Heading {level}")
        _remove_numbering(p)
        run = p.add_run(text)
        run.font.color.rgb = ACCENT_COLOR
        return p

    def _add_separator_line(doc: Document):
        p = doc.add_paragraph()
        _remove_numbering(p)
        p_pr  = p._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), SEPARATOR_COLOR)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _set_cell_background(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _add_company_section(doc: Document, heading: str, section: Dict):
        _add_clean_heading(doc, heading, level=2)
        _add_clean_heading(doc, "포트폴리오 다각화", level=3)
        _add_bullet_items(doc, section.get("portfolio_diversification", []))
        _add_clean_heading(doc, "핵심 경쟁력", level=3)
        _add_bullet_items(doc, section.get("core_competencies", []))
        p = doc.add_paragraph()
        _remove_numbering(p)
        p.add_run(f"전략 방향: {section.get('strategic_direction', '')}")
        watchpoints = section.get("key_watchpoints", [])
        if watchpoints:
            wp = doc.add_paragraph()
            _remove_numbering(wp)
            wp.add_run("Watchpoints")
            _add_bullet_items(doc, watchpoints)

    def _add_swot_table(doc: Document):
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "구분"
        hdr[1].text = "A기업 (SK On)"
        hdr[2].text = "B기업 (CATL)"
        hdr[3].text = "전략적 시사점 (비교 평가)"
        for cell in hdr:
            _set_cell_background(cell, TABLE_HEADER_FILL)
        for row in swot.get("swot_comparison_table", []):
            cells = table.add_row().cells
            cells[0].text = row.get("category", "")
            cells[1].text = row.get("company_a_summary", "")
            cells[2].text = row.get("company_b_summary", "")
            cells[3].text = row.get("strategic_implication", "")
            _set_cell_background(cells[0], TABLE_FIRST_COL_FILL)

    # ── 문서 작성 ──────────────────────────────
    doc = Document()

    title_p = doc.add_paragraph(style="Title")
    _remove_numbering(title_p)
    title_run = title_p.add_run(report.get("title", "배터리 시장 전략 분석 보고서"))
    title_run.font.color.rgb = ACCENT_COLOR

    _add_clean_heading(doc, "목차", level=1)
    _add_bullet_items(doc, [
        "1. Summary",
        "2. 시장 배경(배터리 시장 환경 변화)",
        "3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력",
        "4. Comparative SWOT",
        "5. 종합 시사점",
        "6. Reference",
    ])
    _add_separator_line(doc)

    _add_clean_heading(doc, "1. Summary", level=1)
    doc.add_paragraph(report.get("summary", ""))

    _add_clean_heading(doc, "2. 시장 배경(배터리 시장 환경 변화)", level=1)
    _add_bullet_items(doc, report.get("market_background", []))

    _add_clean_heading(doc, "3. 각 기업별 포트폴리오 다각화 및 핵심 경쟁력", level=1)
    _add_company_section(doc, "3-1. SK On", report.get("sk_on_section", {}))
    _add_company_section(doc, "3-2. CATL",  report.get("catl_section", {}))

    _add_clean_heading(doc, "4. Comparative SWOT", level=1)
    _add_clean_heading(doc, "4-1. S,W,O,T 별 비교 기준 및 주목한 포인트", level=2)
    _add_bullet_items(doc, report.get("comparative_swot_focus_points", []))
    _add_clean_heading(doc, "4-2. S,W,O,T 별 기업 비교", level=2)
    _add_bullet_items(doc, report.get("comparative_swot_company_comparison", []))
    _add_clean_heading(doc, "4-3. SWOT 비교 요약 table", level=2)
    _add_swot_table(doc)

    _add_clean_heading(doc, "5. 종합 시사점", level=1)
    _add_bullet_items(doc, report.get("integrated_implications", []))

    _add_clean_heading(doc, "6. Reference", level=1)
    _add_bullet_items(doc, report.get("references", []))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.node_exit("publish_report", duration_sec=0, status="ok",
                     metadata={"path": str(output_path)})
